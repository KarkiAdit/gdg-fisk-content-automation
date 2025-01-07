import os
import tempfile
import mimetypes
import uuid
from docx import Document

class FileHandler:
    """
    A utility class for handling file operations in a temporary directory, including saving,
    deleting, and cleaning up files.

    This class creates a temporary folder upon initialization and provides methods for:
    - Saving files with specified filenames.
    - Deleting specific files from the temporary folder.
    - Cleaning up (deleting) all files and the temporary folder itself.

    Attributes:
        temp_folder (str): The path to the temporary directory created for file operations.
    """

    def __init__(self):
        """
        Initializes the FileHandler and creates a temporary folder for file operations.

        This method creates a temporary directory using Python's `tempfile` module. All file operations 
        performed by this class will be contained within this directory.
        """
        self.temp_folder = tempfile.mkdtemp()
        print(f"Temporary folder created: {self.temp_folder}")

    def read_file(self, file_path):
        """
        Reads the content of the file at the given path, and determines whether to open it in binary ('rb') or 
        text ('r') mode based on the file's type (e.g., by extension or MIME type).

        Args:
            file_path (str): The path to the file to read.

        Returns:
            str or bytes: The file content as a string (if text) or bytes (if binary).
        """
        try:
            # Determine if the file is binary or text
            mime_type, _ = mimetypes.guess_type(file_path)   
            if mime_type and mime_type.startswith('text'):
                # Open file in text mode
                with open(file_path, 'r', encoding='utf-8') as file:
                    file_content = file.read()
                return file_content
            else:
                # Open file in binary mode
                with open(file_path, 'rb') as file:
                    file_content = file.read()
                return file_content
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            return None
        
    def save_file(self, file_content, output_filename="output.docx"):
        """
        Saves the provided content to a file within the temporary folder.

        Args:
            file_content (str or bytes): The content to save (can be text or binary).
            output_filename (str): The name of the file to save in the temporary folder. Defaults to "output.docx".

        Returns:
            str: The full path to the saved file if successful, None otherwise.
        """
        if not file_content:
            print("Failed to save file: No content provided.")
            return None
        
        file_path = os.path.join(self.temp_folder, output_filename)
        
        try:
            # Determine if the content is text or binary and save accordingly
            if isinstance(file_content, str):
                # Ensure proper newline handling for text files, especially for Windows (if relevant)
                file_content = file_content.replace('\r\n', '\n')  # Normalize line endings to LF (Unix-style)
                # Save as text (use "w" mode for writing text files)
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(file_content)
            elif isinstance(file_content, bytes):
                # Save as binary (use "wb" mode for writing binary files)
                with open(file_path, "wb") as file:
                    file.write(file_content)
            else:
                print("Unsupported file content type")
                return None
            print(f"File saved at {file_path}")
            return file_path
        except Exception as error:
            print(f"Error saving file: {error}")
            return None

    def delete_file(self, filename):
        """
        Deletes a file with the specified name from the temporary folder.

        Args:
            filename (str): The name of the file to delete from the temporary folder.

        Returns:
            bool: True if the file was successfully deleted, False if the file was not found or an error occurred.
        """
        file_path = os.path.join(self.temp_folder, filename)
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"File deleted: {file_path}")
                return True
            else:
                print(f"File not found: {file_path}")
                return False
        except Exception as error:
            print(f"Error deleting file: {error}")
            return False

    def cleanup(self):
        """
        Deletes all files in the temporary folder and removes the folder itself.

        This method iterates through all files in the temporary folder, deleting each one, 
        and then removes the folder. Use this method to clean up resources after file operations are complete.

        Returns:
            None
        """
        try:
            # Delete each file in the temporary folder
            for filename in os.listdir(self.temp_folder):
                file_path = os.path.join(self.temp_folder, filename)
                os.remove(file_path)
            os.rmdir(self.temp_folder)
            print(f"Temporary folder {self.temp_folder} and all contents have been deleted.")
        except Exception as error:
            print(f"Error cleaning up temp folder: {error}")

    def update_file_with_uuid(self, target_path):
        """
        Updates the first paragraph of a .docx file to include a UUID, if not already present.

        This method checks if the first paragraph of the specified .docx file contains a UUID 
        in the format `file_id: "uuid"`. If a valid UUID is already present, it is left unchanged. 
        If the `file_id:` marker is present without a UUID, a new UUID is generated and added. 
        If the `file_id:` marker is not present, a new paragraph is prepended to the file with the UUID.

        Args:
            target_path (str): The path to the .docx file to update.

        Returns:
            tuple:
                - file_uuid (str): The UUID found or generated for the file.
                - added_new_uuid (bool): True if a new UUID was added, False if a valid UUID was already present.

        Raises:
            Exception: If there is an error reading or saving the file.
        """
        try:
            doc = Document(target_path)
            first_paragraph = doc.paragraphs[0].text.strip() if doc.paragraphs else ""
            uuid_marker = "file_id:"
            file_uuid = ""
            added_new_uuid = False
            if first_paragraph.startswith(uuid_marker):
                # Extract the UUID if present
                file_uuid = first_paragraph.split(uuid_marker)[-1].strip().strip('"')
                if file_uuid:
                    print(f"UUID already present: {file_uuid}")
                else:
                    print("UUID placeholder present but no UUID. Generating new UUID...")
                    file_uuid = str(uuid.uuid4())
                    doc.paragraphs[0].text = f'{uuid_marker} "{file_uuid}"'
                    added_new_uuid = True
            else:
                # Add a new UUID at the beginning
                print("UUID not present. Generating and prepending new UUID...")
                file_uuid = str(uuid.uuid4())
                new_paragraph = doc.add_paragraph(f'{uuid_marker} "{file_uuid}"')
                doc.paragraphs.insert(0, new_paragraph)
                added_new_uuid = True
            # Save changes if a new UUID was added
            if added_new_uuid:
                doc.save(target_path)
                print(f"Updated file saved to {target_path}")
            return file_uuid, added_new_uuid
        except Exception as e:
            print(f"Error generating the UUID: {e}")
            raise